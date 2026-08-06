from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "QUESTIONNAIRE_ITEM_BY_ITEM_REVIEW.md"
OUTPUT = ROOT / "QUESTIONNAIRE_ITEM_BY_ITEM_REVIEW.docx"

# compact_reference_guide preset with one named brand-color override.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11
MARGIN_IN = 1.0
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

FONT_LATIN = "Arial Unicode MS"
FONT_CJK = "Arial Unicode MS"
COLOR_BRAND = "0F766E"
COLOR_BRAND_DARK = "12312D"
COLOR_BODY = "172322"
COLOR_MUTED = "53625F"
COLOR_RULE = "CFE0DC"
COLOR_FILL = "E8F3F1"
COLOR_REVIEW = "F7FAF9"
COLOR_WHITE = "FFFFFF"


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, bold=None, color=COLOR_BODY, font=FONT_LATIN):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLOR_BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10, COLOR_BRAND),
        "Heading 2": (13, 14, 7, COLOR_BRAND),
        "Heading 3": (12, 10, 5, COLOR_BRAND_DARK),
    }
    for name, (size, before, after, color) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=9, color=COLOR_MUTED)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("EG BioMed | 問卷逐題審核規格")
    set_run_font(run, size=9, bold=True, color=COLOR_MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("第 ")
    set_run_font(run, size=9, color=COLOR_MUTED)
    add_page_field(p, "PAGE")
    run = p.add_run(" 頁／共 ")
    set_run_font(run, size=9, color=COLOR_MUTED)
    add_page_field(p, "NUMPAGES")
    run = p.add_run(" 頁")
    set_run_font(run, size=9, color=COLOR_MUTED)


def add_inline_markdown(paragraph, text, size=11, color=COLOR_BODY):
    token_pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size - 0.5, color=COLOR_BRAND_DARK, font="Consolas")
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def set_paragraph_border(paragraph, color=COLOR_RULE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_title_block(doc, metadata_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("問卷逐題審核")
    set_run_font(run, size=25, bold=True, color=COLOR_BRAND_DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("文字、選項、跳題與追問邏輯規格")
    set_run_font(run, size=14, color=COLOR_BRAND)
    set_paragraph_border(p, color=COLOR_BRAND, size="12")

    table = doc.add_table(rows=len(metadata_lines) + 1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2600, 6760])
    for index, line in enumerate(metadata_lines):
        label, value = line.split("：", 1)
        set_cell_fill(table.cell(index, 0), COLOR_FILL)
        p1 = table.cell(index, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=10, bold=True, color=COLOR_BRAND_DARK)
        p2 = table.cell(index, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        add_inline_markdown(p2, value.strip(), size=10)

    row = table.rows[-1]
    merged = row.cells[0].merge(row.cells[1])
    set_cell_fill(merged, COLOR_REVIEW)
    p = merged.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("用途：供前端、後端、模型、臨床與法規人員逐題確認。本文件不代表題目已通過最終審查。")
    set_run_font(run, size=10, color=COLOR_MUTED)


def parse_markdown_table(lines, start_index):
    rows = []
    index = start_index
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip()[1:-1]
        protected = raw.replace("\\|", "\u0000")
        cells = [cell.strip().replace("\u0000", "|") for cell in protected.split("|")]
        if not all(re.fullmatch(r"-+", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_options_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    widths = [2300, 3300, 3760]
    for row_index, source_row in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for column_index, value in enumerate(source_row[:3]):
            cell = row.cells[column_index]
            if row_index == 0:
                set_cell_fill(cell, COLOR_BRAND)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            clean = value.strip("`")
            run = p.add_run(clean)
            set_run_font(
                run,
                size=9 if row_index else 9.5,
                bold=row_index == 0,
                color=COLOR_WHITE if row_index == 0 else COLOR_BODY,
                font="Consolas" if column_index == 0 and row_index else FONT_LATIN,
            )
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_review_box(doc):
    p = doc.add_paragraph(style="Heading 3")
    p.add_run("審核結果")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [4680, 4680])
    choices = ["[ ] 保留", "[ ] 修改", "[ ] 刪除", "[ ] 待臨床／模型／法規確認"]
    for index, choice in enumerate(choices):
        cell = table.cell(index // 2, index % 2)
        set_cell_fill(cell, COLOR_REVIEW)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(choice)
        set_run_font(run, size=10, color=COLOR_BODY)
    merged = table.rows[2].cells[0].merge(table.rows[2].cells[1])
    set_cell_fill(merged, COLOR_WHITE)
    p = merged.paragraphs[0]
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("修改說明：")
    set_run_font(run, size=10, bold=True, color=COLOR_MUTED)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def split_source(source):
    lines = source.splitlines()
    title = lines[0].lstrip("# ").strip()
    common_index = lines.index("## 共通規則")
    review_index = lines.index("## 逐題審核")
    metadata = [line.strip() for line in lines[1:common_index] if line.strip()]
    common = [line.strip() for line in lines[common_index + 1:review_index] if line.strip()]
    question_lines = lines[review_index + 1:]

    starts = [index for index, line in enumerate(question_lines) if re.match(r"^## \d+\. ", line)]
    blocks = []
    for position, start in enumerate(starts):
        end = starts[position + 1] if position + 1 < len(starts) else len(question_lines)
        blocks.append(question_lines[start:end])
    return title, metadata, common, blocks


def extract_module(block):
    for line in block:
        match = re.match(r"^- \*\*段落\*\*：(.+)$", line.strip())
        if match:
            return match.group(1).strip()
    return "其他"


def build_docx():
    source = SOURCE.read_text(encoding="utf-8")
    _, metadata, common, blocks = split_source(source)

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    doc.core_properties.title = "問卷逐題審核文字與邏輯規格"
    doc.core_properties.subject = "v19.4 題庫、跳題與追問規格"
    doc.core_properties.author = "EG BioMed"
    doc.core_properties.keywords = "questionnaire, review, v19.4, EG BioMed"

    add_title_block(doc, metadata)

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("共通規則")
    for item in common:
        match = re.match(r"^(\d+)\.\s+(.+)$", item)
        if not match:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        number = p.add_run(f"{match.group(1)}. ")
        set_run_font(number, size=11, bold=True, color=COLOR_BRAND)
        add_inline_markdown(p, match.group(2), size=11)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("逐題審核")
    set_run_font(run, size=20, bold=True, color=COLOR_BRAND_DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("共 77 題，排除知情同意。題號為本審核文件的順序；「原始總題序」含知情同意。")
    set_run_font(run, size=11, color=COLOR_MUTED)

    current_module = None
    for block_index, block in enumerate(blocks):
        module = extract_module(block)
        if module != current_module:
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(module)
            p.paragraph_format.space_before = Pt(4)
            current_module = module

        heading = block[0].lstrip("# ").strip()
        p = doc.add_paragraph(style="Heading 2")
        p.add_run(heading)

        index = 1
        while index < len(block):
            line = block[index].strip()
            if not line or line == "---":
                index += 1
                continue
            if line == "**選項與固定代碼**":
                caption = doc.add_paragraph()
                caption.paragraph_format.space_before = Pt(5)
                caption.paragraph_format.space_after = Pt(4)
                run = caption.add_run("選項與固定代碼")
                set_run_font(run, size=10.5, bold=True, color=COLOR_BRAND_DARK)
                index += 1
                while index < len(block) and not block[index].strip():
                    index += 1
                rows, index = parse_markdown_table(block, index)
                add_options_table(doc, rows)
                continue
            if line == "**審核結果**":
                add_review_box(doc)
                while index < len(block) and block[index].strip() != "---":
                    index += 1
                continue
            if line.startswith("- **"):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                add_inline_markdown(p, line[2:], size=10.5)
            index += 1

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
